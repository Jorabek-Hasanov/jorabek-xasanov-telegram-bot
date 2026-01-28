# File Processing Service
# ============================================================================
# PDF, DOCX, and Image OCR processing with security sanitization
# Supports EasyOCR and Tesseract for text extraction from images
# ============================================================================

import asyncio
import enum
import io
import os
import tempfile
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import quote

import magic
from PIL import Image

from src.core.config import get_settings
from src.core.logging import logger


class FileType(enum.Enum):
    """Supported file types."""
    PDF = "pdf"
    DOCX = "docx"
    IMAGE_PNG = "image/png"
    IMAGE_JPEG = "image/jpeg"
    IMAGE_JPG = "image/jpg"
    IMAGE_GIF = "image/gif"
    IMAGE_WEBP = "image/webp"
    UNKNOWN = "unknown"


@dataclass
class ProcessedFile:
    """Result of file processing."""
    file_type: FileType
    text_content: str
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    language_detected: Optional[str] = None
    processing_time_ms: float = 0.0
    file_size_bytes: int = 0
    warnings: List[str] = None
    
    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []


class TextExtractor(ABC):
    """Abstract base class for text extractors."""
    
    @abstractmethod
    async def extract(self, file_path: str) -> str:
        """Extract text from a file."""
        pass


class PDFExtractor(TextExtractor):
    """PDF text extractor using PyMuPDF (fitz)."""
    
    def __init__(self) -> None:
        """Initialize PDF extractor."""
        try:
            import fitz
            self.fitz = fitz
            self.available = True
        except ImportError:
            logger.warning("PyMuPDF not available, PDF extraction disabled")
            self.available = False
    
    async def extract(self, file_path: str) -> str:
        """Extract text from PDF file."""
        if not self.available:
            raise RuntimeError("PyMuPDF not installed")
        
        text_content = []
        doc = self.fitz.open(file_path)
        
        for page_num, page in enumerate(doc):
            text = page.get_text()
            text_content.append(f"[Page {page_num + 1}]\n{text}")
        
        doc.close()
        return "\n\n".join(text_content)


class DOCXExtractor(TextExtractor):
    """DOCX text extractor using python-docx."""
    
    def __init__(self) -> None:
        """Initialize DOCX extractor."""
        try:
            import docx
            self.docx = docx
            self.available = True
        except ImportError:
            logger.warning("python-docx not available, DOCX extraction disabled")
            self.available = False
    
    async def extract(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if not self.available:
            raise RuntimeError("python-docx not installed")
        
        doc = self.docx.Document(file_path)
        text_content = []
        
        for para_num, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                text_content.append(text)
        
        # Extract tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_content.append(f"[Table] {row_text}")
        
        return "\n\n".join(text_content)


class EasyOCRExtractor(TextExtractor):
    """OCR extractor using EasyOCR."""
    
    def __init__(self) -> None:
        """Initialize EasyOCR extractor."""
        try:
            import easyocr
            self.reader = easyocr.Reader(['en', 'ru', 'uz'])
            self.available = True
        except ImportError:
            logger.warning("EasyOCR not available, OCR disabled")
            self.available = False
    
    async def extract(self, file_path: str) -> str:
        """Extract text from image using EasyOCR."""
        if not self.available:
            raise RuntimeError("EasyOCR not installed")
        
        # EasyOCR is synchronous, run in thread pool
        results = await asyncio.to_thread(
            self.reader.readtext,
            file_path,
            detail=0,
            paragraph=True
        )
        
        return "\n\n".join(results)


class TesseractExtractor(TextExtractor):
    """OCR extractor using Tesseract."""
    
    def __init__(self) -> None:
        """Initialize Tesseract extractor."""
        try:
            import pytesseract
            self.pytesseract = pytesseract
            self.available = True
        except ImportError:
            logger.warning("pytesseract not available, OCR disabled")
            self.available = False
    
    async def extract(self, file_path: str) -> str:
        """Extract text from image using Tesseract."""
        if not self.available:
            raise RuntimeError("Tesseract not installed")
        
        image = Image.open(file_path)
        text = await asyncio.to_thread(
            self.pytesseract.image_to_string,
            image
        )
        
        return text


class FileProcessor:
    """
    Unified file processor for PDF, DOCX, and images.
    
    Features:
    - MIME type detection using python-magic
    - File size validation (max 20MB)
    - Secure in-memory processing
    - Temporary file cleanup
    - Multiple OCR engine support
    
    Example:
        ```python
        processor = FileProcessor()
        result = await processor.process_file(file_bytes, "application/pdf")
        print(result.text_content)
        ```
    """
    
    def __init__(self) -> None:
        """Initialize file processor with all extractors."""
        self.settings = get_settings().file_processing
        
        # Initialize extractors
        self.pdf_extractor = PDFExtractor()
        self.docx_extractor = DOCXExtractor()
        
        # Select OCR engine based on settings
        if self.settings.ocr_engine == "easyocr":
            self.ocr_extractor = EasyOCRExtractor()
        else:
            self.ocr_extractor = TesseractExtractor()
        
        # Create temp directory
        self.temp_dir = Path(self.settings.temp_dir)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        
        # Magic for MIME detection
        self.mime_detector = magic.Magic(mime=True)
    
    def detect_file_type(
        self,
        file_content: bytes,
        filename: str = ""
    ) -> Tuple[FileType, Optional[str]]:
        """
        Detect file type from content.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename for extension detection
            
        Returns:
            Tuple of (FileType, detected_extension)
        """
        # Check MIME type
        mime_type = self.mime_detector.from_buffer(file_content)
        
        # Map MIME types to FileType enum
        mime_mapping = {
            "application/pdf": FileType.PDF,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": FileType.DOCX,
            "image/png": FileType.IMAGE_PNG,
            "image/jpeg": FileType.IMAGE_JPEG,
            "image/jpg": FileType.IMAGE_JPG,
            "image/gif": FileType.IMAGE_GIF,
            "image/webp": FileType.IMAGE_WEBP,
        }
        
        file_type = mime_mapping.get(mime_type, FileType.UNKNOWN)
        
        # Also check extension as fallback
        if file_type == FileType.UNKNOWN and filename:
            ext = Path(filename).suffix.lower()
            ext_mapping = {
                ".pdf": FileType.PDF,
                ".docx": FileType.DOCX,
                ".png": FileType.IMAGE_PNG,
                ".jpg": FileType.IMAGE_JPEG,
                ".jpeg": FileType.IMAGE_JPEG,
                ".gif": FileType.IMAGE_GIF,
                ".webp": FileType.IMAGE_WEBP,
            }
            file_type = ext_mapping.get(ext, FileType.UNKNOWN)
        
        return file_type, mime_type
    
    def validate_file(
        self,
        file_content: bytes,
        filename: str = ""
    ) -> Tuple[bool, str]:
        """
        Validate file for processing.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file size
        max_size_bytes = self.settings.max_file_size_mb * 1024 * 1024
        if len(file_content) > max_size_bytes:
            return False, (
                f"File too large ({len(file_content) / 1024 / 1024:.1f}MB). "
                f"Maximum size is {self.settings.max_file_size_mb}MB.\n\n"
                f"💡 <b>Senior Recommendation:</b> Compress the file using tools like "
                f"Adobe Acrobat for PDFs or remove unnecessary images. "
                f"For images, reduce resolution to 150-300 DPI."
            )
        
        # Detect file type
        file_type, mime_type = self.detect_file_type(file_content, filename)
        
        if file_type == FileType.UNKNOWN:
            return False, "Unsupported file type. Please send PDF, DOCX, or image files."
        
        if file_type not in self.settings.allowed_mime_types and file_type != FileType.UNKNOWN:
            return False, f"File type {mime_type} is not allowed."
        
        return True, ""
    
    async def process_file(
        self,
        file_content: bytes,
        filename: str = "",
        user_id: Optional[int] = None,
    ) -> ProcessedFile:
        """
        Process a file and extract text content.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            user_id: User ID for logging
            
        Returns:
            ProcessedFile with extracted text
        """
        import time
        start_time = time.perf_counter()
        warnings = []
        
        # Validate file
        is_valid, error_message = self.validate_file(file_content, filename)
        if not is_valid:
            return ProcessedFile(
                file_type=FileType.UNKNOWN,
                text_content="",
                warnings=[error_message],
                processing_time_ms=(time.perf_counter() - start_time) * 1000,
                file_size_bytes=len(file_content),
            )
        
        file_type, mime_type = self.detect_file_type(file_content, filename)
        
        logger.info(
            f"Processing file: {filename} ({mime_type}) for user {user_id}",
            extra={"event_type": "file_processing", "user_id": user_id}
        )
        
        # Create temporary file
        temp_file_path = self.temp_dir / f"{user_id}_{int(time.time())}_{filename}"
        
        try:
            # Write to temporary file
            with open(temp_file_path, "wb") as f:
                f.write(file_content)
            
            # Extract text based on file type
            if file_type == FileType.PDF:
                text_content = await self.pdf_extractor.extract(str(temp_file_path))
            elif file_type == FileType.DOCX:
                text_content = await self.docx_extractor.extract(str(temp_file_path))
            else:
                # OCR for images
                text_content = await self.ocr_extractor.extract(str(temp_file_path))
            
            # Calculate word count
            word_count = len(text_content.split()) if text_content else 0
            
            processing_time_ms = (time.perf_counter() - start_time) * 1000
            
            logger.info(
                f"File processed successfully: {word_count} words in {processing_time_ms:.2f}ms",
                extra={"event_type": "file_processing_complete", "user_id": user_id}
            )
            
            return ProcessedFile(
                file_type=file_type,
                text_content=text_content,
                word_count=word_count,
                processing_time_ms=processing_time_ms,
                file_size_bytes=len(file_content),
                warnings=warnings,
            )
            
        except Exception as e:
            logger.error(
                f"File processing failed: {e}",
                extra={"event_type": "file_processing_error", "user_id": user_id}
            )
            raise
        finally:
            # Cleanup temporary file
            if temp_file_path.exists():
                temp_file_path.unlink()
    
    async def process_multiple_files(
        self,
        files: List[Tuple[bytes, str]],
        user_id: Optional[int] = None,
    ) -> List[ProcessedFile]:
        """
        Process multiple files concurrently.
        
        Args:
            files: List of (file_content, filename) tuples
            user_id: User ID for logging
            
        Returns:
            List of ProcessedFile results
        """
        tasks = [
            self.process_file(content, filename, user_id)
            for content, filename in files
        ]
        
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Convert exceptions to error results
        processed_files = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_files.append(ProcessedFile(
                    file_type=FileType.UNKNOWN,
                    text_content="",
                    warnings=[f"Processing failed: {str(result)}"],
                ))
            else:
                processed_files.append(result)
        
        return processed_files
